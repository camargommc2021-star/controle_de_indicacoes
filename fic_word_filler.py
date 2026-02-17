"""
FIC Word Filler v2 - Preenche template Word do FIC com layout EXATO do modelo

Template oficial do CRCEA-SE
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_UNDERLINE, WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime
import re
import json


class FICWordFiller:
    """Preenche o template Word do FIC com dados do sistema - Layout EXATO"""
    
    def __init__(self):
        self.template_path = "assets/FIC_template.docx"
        self.mapeamento_funcoes = self._carregar_mapeamento_funcoes()
        
    def _carregar_mapeamento_funcoes(self):
        """Carrega o mapeamento de habilitações para funções"""
        try:
            with open('data/mapeamento_funcoes.json', 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return {
                "S": "SUPERVISOR",
                "I": "INSTRUTOR", 
                "O": "OPERADOR",
                "F": "FMC",
                "CHEQ": "CHEFE DE EQUIPE",
                "E": "ESTAGIÁRIO"
            }
    
    def _get_funcao_completa(self, habilitacao):
        """Retorna o nome completo da função baseado na habilitação"""
        if not habilitacao:
            return ""
        hab_upper = str(habilitacao).strip().upper()
        return self.mapeamento_funcoes.get(hab_upper, hab_upper)
    
    def _to_upper(self, valor):
        """Converte valor para maiúsculo, tratando valores None"""
        if valor is None:
            return ''
        return str(valor).strip().upper()
    
    def _formatar_cpf(self, cpf):
        """Formata CPF no padrão XXX.XXX.XXX-XX"""
        if not cpf:
            return ''
        cpf_limpo = re.sub(r'\D', '', str(cpf))
        if len(cpf_limpo) != 11:
            return self._to_upper(cpf)
        return f"{cpf_limpo[:3]}.{cpf_limpo[3:6]}.{cpf_limpo[6:9]}-{cpf_limpo[9:]}"
    
    def _formatar_saram(self, saram):
        """Formata SARAM removendo caracteres não numéricos"""
        if not saram:
            return ''
        saram_limpo = re.sub(r'\D', '', str(saram))
        return saram_limpo
    
    def _calcular_tempo_servico(self, data_praca):
        """Calcula tempo de serviço em anos e meses a partir da data de praça"""
        if not data_praca:
            return '', ''
        
        try:
            formatos = ['%d/%m/%Y', '%Y-%m-%d', '%d-%m-%Y', '%m/%d/%Y']
            data_praca_dt = None
            
            for formato in formatos:
                try:
                    data_praca_dt = datetime.strptime(str(data_praca).strip(), formato)
                    break
                except ValueError:
                    continue
            
            if not data_praca_dt:
                return '', ''
            
            hoje = datetime.now()
            anos = hoje.year - data_praca_dt.year
            meses = hoje.month - data_praca_dt.month
            
            if meses < 0:
                anos -= 1
                meses += 12
            
            return str(anos), str(meses)
        except:
            return '', ''
    
    def preencher_fic(self, dados_fic, output_path=None):
        """
        Preenche o template FIC com os dados fornecidos - Layout EXATO do modelo
        """
        # Abrir template
        doc = Document(self.template_path)
        
        # O template usa uma tabela principal
        if doc.tables:
            self._preencher_tabela_exato(doc.tables[0], dados_fic)
        
        # Salvar em buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Se tiver output_path, salvar também no disco
        if output_path:
            doc.save(output_path)
        
        return buffer
    
    def _preencher_tabela_exato(self, table, dados):
        """Preenche a tabela replicando EXATAMENTE o layout do modelo"""
        anos_servico, meses_servico = self._calcular_tempo_servico(dados.get('Data_Praca'))
        
        # Processar cada linha específica conforme o modelo
        for row_idx, row in enumerate(table.rows):
            if row_idx == 0:
                # Linha 0: Cabeçalho - não alterar
                pass
            elif row_idx == 1:
                # Linha 1: Instrução - não alterar
                pass
            elif row_idx == 2:
                # Linha 2: CÓDIGO DO CURSO, TURMA, LOCAL DO CURSO/GT, COMANDO
                self._preencher_linha_2(row, dados)
            elif row_idx == 3:
                # Linha 3: NOME DO CURSO
                self._preencher_linha_3(row, dados)
            elif row_idx == 4:
                # Linha 4: DATAS PRESENCIAL
                self._preencher_linha_4(row, dados)
            elif row_idx == 5:
                # Linha 5: DATAS EAD
                self._preencher_linha_5(row, dados)
            elif row_idx == 6:
                # Linha 6: PDP
                self._preencher_linha_6(row, dados)
            elif row_idx == 7:
                # Linha 7: POSTO/GRAD/ESP/NOME COMPLETO
                self._preencher_linha_7(row, dados)
            elif row_idx == 8:
                # Linha 8: OM DO INDICADO, CPF, SARAM
                self._preencher_linha_8(row, dados)
            elif row_idx == 9:
                # Linha 9: E-MAIL FUNCIONAL, TELEFONE
                self._preencher_linha_9(row, dados)
            elif row_idx == 10:
                # Linha 10: FUNÇÃO ATUAL, DATA ÚLTIMA PROMOÇÃO
                self._preencher_linha_10(row, dados)
            elif row_idx == 11:
                # Linha 11: FUNÇÃO QUE EXERCERÁ APÓS O CURSO
                self._preencher_linha_11(row, dados)
            elif row_idx == 12:
                # Linha 12: TEMPO DE SERVIÇO, PRÉ-REQUISITOS
                self._preencher_linha_12(row, dados, anos_servico, meses_servico)
            elif row_idx == 13:
                # Linha 13: Cabeçalho campos do chefe
                pass
            elif row_idx == 14:
                # Linha 14: JUSTIFICATIVA
                self._preencher_linha_14(row, dados)
            elif row_idx in [15, 16, 17, 18, 19]:
                # Linhas 15-19: Questionários SIM/NÃO
                self._preencher_questionario(row, dados, row_idx)
            elif row_idx == 20:
                # Linha 20: Assinatura Chefe Imediato
                self._preencher_assinatura(row, dados, 'chefe')
            elif row_idx == 21:
                # Linha 21: Assinatura Responsável
                self._preencher_assinatura(row, dados, 'resp')
    
    def _preencher_linha_2(self, row, dados):
        """Linha 2: CÓDIGO DO CURSO, TURMA, LOCAL DO CURSO/GT, COMANDO"""
        codigo = self._to_upper(dados.get('Codigo_Curso', ''))
        turma = self._to_upper(dados.get('Turma', ''))
        local = self._to_upper(dados.get('Local_GT', ''))
        comando = self._to_upper(dados.get('Comando', ''))
        
        for cell in row.cells:
            text = cell.text.strip()
            if 'CÓDIGO DO CURSO:' in text:
                self._substituir_em_cell(cell, 'CÓDIGO DO CURSO:', codigo)
            elif 'TURMA:' in text and turma:
                self._substituir_em_cell(cell, 'TURMA:', turma)
            elif 'LOCAL DO CURSO/GT:' in text:
                self._substituir_em_cell(cell, 'LOCAL DO CURSO/GT:', local)
            elif 'COMANDO:' in text:
                self._substituir_em_cell(cell, 'COMANDO:', comando)
    
    def _preencher_linha_3(self, row, dados):
        """Linha 3: NOME DO CURSO"""
        nome = self._to_upper(dados.get('Nome_Curso', ''))
        for cell in row.cells:
            if 'NOME DO CURSO:' in cell.text:
                self._substituir_em_cell(cell, 'NOME DO CURSO:', nome)
    
    def _preencher_linha_4(self, row, dados):
        """Linha 4: DATAS PRESENCIAL"""
        inicio = self._to_upper(dados.get('Data_Inicio_Presencial', ''))
        termino = self._to_upper(dados.get('Data_Termino_Presencial', ''))
        
        for cell in row.cells:
            text = cell.text.strip()
            if 'DATA DE INÍCIO (Presencial):' in text:
                self._substituir_texto_inline(cell, 'DATA DE INÍCIO (Presencial):', inicio)
            elif 'DATA DE TÉRMINO (Presencial):' in text:
                self._substituir_texto_inline(cell, 'DATA DE TÉRMINO (Presencial):', termino)
    
    def _preencher_linha_5(self, row, dados):
        """Linha 5: DATAS EAD"""
        inicio = self._to_upper(dados.get('Data_Inicio_Distancia', ''))
        termino = self._to_upper(dados.get('Data_Termino_Distancia', ''))
        
        for cell in row.cells:
            text = cell.text.strip()
            if 'DATA DE INÍCIO (A distância):' in text:
                self._substituir_texto_inline(cell, 'DATA DE INÍCIO (A distância):', inicio)
            elif 'DATA DE TÉRMINO (A distância):' in text:
                self._substituir_texto_inline(cell, 'DATA DE TÉRMINO (A distância):', termino)
    
    def _preencher_linha_6(self, row, dados):
        """Linha 6: PDP"""
        ppd = self._to_upper(dados.get('PPD_Civil', ''))
        sim_marcado = '( X ) SIM' if ppd == 'SIM' else '(   ) SIM'
        nao_marcado = '( X ) NÃO' if ppd == 'NÃO' or ppd == '' else '(   ) NÃO'
        
        for cell in row.cells:
            if 'PREVISTO NA PDP' in cell.text:
                for paragraph in cell.paragraphs:
                    paragraph.clear()
                    p = cell.paragraphs[0]
                    p.add_run("EM CASO DE ALUNO CIVIL, O CURSO ESTÁ PREVISTO NA PDP?").font.size = Pt(9)
                    p.add_run(f"    {sim_marcado}        {nao_marcado}").font.size = Pt(9)
    
    def _preencher_linha_7(self, row, dados):
        """Linha 7: POSTO/GRAD/ESP/NOME COMPLETO"""
        posto = self._to_upper(dados.get('Posto_Graduacao', ''))
        esp = self._to_upper(dados.get('Especialidade', ''))
        nome = self._to_upper(dados.get('Nome_Completo', ''))
        nome_guerra = self._to_upper(dados.get('Nome_Guerra', ''))
        
        for cell in row.cells:
            if 'POSTO/GRAD/ESP/NOME COMPLETO' in cell.text:
                self._preencher_nome_com_sublinhado(cell, posto, esp, nome, nome_guerra)
    
    def _preencher_linha_8(self, row, dados):
        """Linha 8: OM DO INDICADO, CPF, SARAM"""
        om = self._to_upper(dados.get('OM_Indicado', ''))
        cpf = self._formatar_cpf(dados.get('CPF', ''))
        saram = self._formatar_saram(dados.get('SARAM', ''))
        
        for cell in row.cells:
            text = cell.text.strip()
            if 'OM DO INDICADO:' in text:
                self._substituir_em_cell(cell, 'OM DO INDICADO:', om)
            elif 'CPF:' in text and 'SARAM' not in text:
                self._substituir_em_cell(cell, 'CPF:', cpf)
            elif 'SARAM:' in text:
                self._substituir_em_cell(cell, 'SARAM:', saram)
    
    def _preencher_linha_9(self, row, dados):
        """Linha 9: E-MAIL FUNCIONAL, TELEFONE"""
        email = self._to_upper(dados.get('Email', ''))
        telefone = self._to_upper(dados.get('Telefone', ''))
        
        for cell in row.cells:
            text = cell.text.strip()
            if 'E-MAIL' in text or 'EMAIL' in text:
                self._substituir_em_cell(cell, 'E-MAIL FUNCIONAL:', email)
            elif 'TELEFONE:' in text:
                self._substituir_em_cell(cell, 'TELEFONE:', telefone)
    
    def _preencher_linha_10(self, row, dados):
        """Linha 10: FUNÇÃO ATUAL, DATA ÚLTIMA PROMOÇÃO"""
        funcao = self._to_upper(self._get_funcao_completa(dados.get('Funcao_Atual', '')))
        data_prom = self._to_upper(dados.get('Data_Ultima_Promocao', ''))
        
        for cell in row.cells:
            text = cell.text.strip()
            if 'FUNÇÃO ATUAL:' in text:
                self._substituir_em_cell(cell, 'FUNÇÃO ATUAL:', funcao)
            elif 'DATA ÚLTIMA PROMOÇÃO:' in text:
                self._substituir_em_cell(cell, 'DATA ÚLTIMA PROMOÇÃO:', data_prom)
    
    def _preencher_linha_11(self, row, dados):
        """Linha 11: FUNÇÃO QUE EXERCERÁ APÓS O CURSO"""
        funcao = self._to_upper(self._get_funcao_completa(dados.get('Funcao_Apos_Curso', '')))
        
        for cell in row.cells:
            if 'FUNÇÃO QUE O INDICADO EXERCERÁ' in cell.text:
                self._substituir_em_cell(cell, 'FUNÇÃO QUE O INDICADO EXERCERÁ APÓS O CURSO/ESTÁGIO:', funcao)
    
    def _preencher_linha_12(self, row, dados, anos, meses):
        """Linha 12: TEMPO DE SERVIÇO"""
        tempo = f"{anos} ANOS E {meses} MESES" if anos and meses else ""
        
        for cell in row.cells:
            text = cell.text.strip()
            if anos and meses and ('____ ANOS' in text or ('ANOS' in text and 'MESES' in text)):
                for paragraph in cell.paragraphs:
                    paragraph.clear()
                    if cell.paragraphs:
                        run = cell.paragraphs[0].add_run(tempo)
                        run.font.size = Pt(9)
    
    def _preencher_linha_14(self, row, dados):
        """Linha 14: JUSTIFICATIVA DO CHEFE"""
        justificativa = self._to_upper(dados.get('Justificativa_Chefe', ''))
        
        for cell in row.cells:
            if 'JUSTIFICATIVA DO CHEFE IMEDIATO:' in cell.text:
                self._substituir_em_cell(cell, 'JUSTIFICATIVA DO CHEFE IMEDIATO:', justificativa)
    
    def _preencher_questionario(self, row, dados, row_idx):
        """Preenche as linhas de questionário SIM/NÃO (15-19)"""
        mapeamento = {
            15: ('Curso_Mapeado', 'O CURSO SOLICITADO ESTÁ MAPEADO'),
            16: ('Progressao_Carreira', 'O CURSO SOLICITADO FAZ PARTE DA PROGRESSÃO'),
            17: ('Comunicado_Indicado', 'O INDICADO FOI COMUNICADO'),
            18: ('Curso_Anterior', 'JÁ REALIZOU O CURSO ANTERIORMENTE'),
            19: ('Ciencia_Dedicacao_EAD', 'O CHEFE ESTÁ CIENTE'),
        }
        
        if row_idx not in mapeamento:
            return
            
        campo, texto_busca = mapeamento[row_idx]
        resposta = self._to_upper(dados.get(campo, 'SIM' if campo != 'Curso_Anterior' else 'NÃO'))
        
        for cell_idx, cell in enumerate(row.cells):
            text = cell.text.strip()
            
            # Verifica se é célula de SIM/NÃO (últimas 4 colunas geralmente)
            if cell_idx >= 9 and '(    )' in text:
                self._marcar_sim_nao_cell(cell, resposta)
            elif 'ANO:' in text and campo == 'Curso_Anterior' and cell_idx <= 8:
                # Preenche o ano se houver
                ano = dados.get('Ano_Curso_Anterior', '')
                if ano and str(ano).strip() and resposta == 'SIM':
                    self._substituir_ano_cell(cell, str(ano))
    
    def _preencher_assinatura(self, row, dados, tipo):
        """Preenche as linhas de assinatura (20 e 21)"""
        if tipo == 'chefe':
            nome = self._to_upper(dados.get('Nome_Chefe_COP', ''))
            posto = self._to_upper(dados.get('Posto_Chefe_COP', ''))
            setor = self._to_upper(dados.get('Setor_Chefe_COP', ''))
            prefixo = 'CHEFE DO'
        else:
            nome = self._to_upper(dados.get('Nome_Responsavel_DACTA', ''))
            posto = self._to_upper(dados.get('Posto_Responsavel_DACTA', ''))
            setor = self._to_upper(dados.get('Setor_Responsavel_DACTA', ''))
            prefixo = 'CHEFE DA'
        
        if not nome:
            return
        
        nome_completo = f"{nome} {posto}".strip()
        local_chefia = f"{prefixo} {setor}" if setor else ""
        
        for cell in row.cells:
            text = cell.text.strip()
            # Procura células que contenham o template de assinatura
            if '_______________________' in text or 'Preencher' in text or 'Posto e Nome' in text:
                self._adicionar_nome_assinatura_completa(cell, nome_completo, local_chefia)
    
    def _preencher_nome_com_sublinhado(self, cell, posto, esp, nome, nome_guerra):
        """Preenche o nome com o nome de guerra sublinhado"""
        for paragraph in cell.paragraphs:
            paragraph.clear()
        
        if cell.paragraphs:
            p = cell.paragraphs[0]
            p.add_run("POSTO/GRAD/ESP/NOME COMPLETO (Sublinhar nome de guerra): ").font.size = Pt(9)
            
            # Monta o nome completo
            prefixo = f"{posto} {esp} ".strip()
            
            if nome_guerra and nome_guerra in nome:
                idx = nome.rfind(nome_guerra)
                if idx >= 0:
                    antes = nome[:idx]
                    depois = nome[idx + len(nome_guerra):]
                    
                    if prefixo:
                        p.add_run(prefixo + " ").font.size = Pt(9)
                    if antes:
                        p.add_run(antes).font.size = Pt(9)
                    
                    # Nome de guerra sublinhado
                    run_guerra = p.add_run(nome_guerra)
                    run_guerra.font.size = Pt(9)
                    run_guerra.underline = WD_UNDERLINE.SINGLE
                    run_guerra.bold = True
                    
                    if depois:
                        p.add_run(depois).font.size = Pt(9)
                else:
                    texto = f"{prefixo} {nome}".strip()
                    p.add_run(texto).font.size = Pt(9)
            else:
                texto = f"{prefixo} {nome}".strip()
                p.add_run(texto).font.size = Pt(9)
    
    def _adicionar_nome_assinatura_completa(self, cell, nome_completo, local_chefia):
        """Adiciona linha de assinatura, nome + posto e função abaixo"""
        # Limpa todos os parágrafos
        for paragraph in cell.paragraphs:
            paragraph.clear()
        
        # Remove todos os parágrafos exceto o primeiro
        while len(cell.paragraphs) > 1:
            p = cell.paragraphs[-1]
            p._element.getparent().remove(p._element)
        
        # Parágrafo 1: Linha de assinatura (centralizada)
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_linha = p1.add_run("___________________________________________________")
        run_linha.font.size = Pt(9)
        
        # Parágrafo 2: Nome do chefe (centralizado, negrito)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_nome = p2.add_run(nome_completo)
        run_nome.font.size = Pt(9)
        run_nome.bold = True
        
        # Parágrafo 3: Função/Local da chefia (centralizado)
        if local_chefia:
            p3 = cell.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_local = p3.add_run(local_chefia)
            run_local.font.size = Pt(9)
    
    def _substituir_em_cell(self, cell, label, novo_valor):
        """Substitui texto em uma célula mantendo o label"""
        for paragraph in cell.paragraphs:
            full_text = paragraph.text
            if label in full_text:
                novo_texto = f"{label} {novo_valor}"
                paragraph.clear()
                run = paragraph.add_run(novo_texto)
                run.font.size = Pt(9)
                return
    
    def _substituir_texto_inline(self, cell, label, novo_valor):
        """Substitui texto mantendo na mesma linha"""
        for paragraph in cell.paragraphs:
            full_text = paragraph.text
            if label in full_text:
                paragraph.clear()
                run = paragraph.add_run(f"{label} {novo_valor}")
                run.font.size = Pt(9)
                run.bold = True
                return
    
    def _marcar_sim_nao_cell(self, cell, resposta):
        """Marca SIM ou NÃO em uma célula"""
        text = cell.text
        
        if not resposta:
            resposta = 'NÃO'
        
        # Marca SIM
        if resposta == 'SIM':
            text = text.replace('(    ) SIM', '( X ) SIM')
            text = text.replace('(   ) SIM', '( X ) SIM')
            text = text.replace('(    ) NÃO', '(   ) NÃO')
            text = text.replace('(    ) N�O', '(   ) NÃO')
        else:  # NÃO
            text = text.replace('(    ) SIM', '(   ) SIM')
            text = text.replace('(   ) SIM', '(   ) SIM')
            text = text.replace('(    ) NÃO', '( X ) NÃO')
            text = text.replace('(    ) N�O', '( X ) NÃO')
        
        # Atualiza célula
        for paragraph in cell.paragraphs:
            paragraph.clear()
        if cell.paragraphs:
            cell.paragraphs[0].add_run(text)
    
    def _substituir_ano_cell(self, cell, ano):
        """Substitui o ano na célula"""
        text = cell.text
        if 'ANO:' in text:
            text = text.replace('_______', ano)
            text = text.replace('______', ano)
            text = text.replace('_____', ano)
            text = text.replace('____', ano)
            text = text.replace('___', ano)
            for paragraph in cell.paragraphs:
                paragraph.clear()
            if cell.paragraphs:
                cell.paragraphs[0].add_run(text)


if __name__ == "__main__":
    # Teste
    dados_teste = {
        'Codigo_Curso': 'SEC001E',
        'Nome_Curso': 'ATC AVSEC',
        'Turma': '01/26',
        'Local_GT': 'ICEA',
        'Comando': 'DECEA',
        'Data_Inicio_Presencial': '05/02/2026',
        'Data_Termino_Presencial': '10/02/2026',
        'Data_Inicio_Distancia': '27/02/2026',
        'Data_Termino_Distancia': '27/02/2026',
        'Posto_Graduacao': '1S',
        'Especialidade': 'BCT',
        'Nome_Completo': 'MAURICIO MENDONCA DE CAMARGO',
        'Nome_Guerra': 'CAMARGO',
        'OM_Indicado': 'CRCEA-SE',
        'CPF': '40309619831',
        'SARAM': '4237447',
        'Email': 'camargommc@fab.mil.br',
        'Telefone': '(11) 2112-3421',
        'Funcao_Atual': 'SUPERVISOR',
        'Data_Ultima_Promocao': '01/01/2020',
        'Funcao_Apos_Curso': 'SUPERVISOR',
        'Data_Praca': '01/01/2009',
        'Pre_Requisitos': 'SIM',
        'Curso_Mapeado': 'SIM',
        'Progressao_Carreira': 'SIM',
        'Comunicado_Indicado': 'SIM',
        'Curso_Anterior': 'NÃO',
        'Ciencia_Dedicacao_EAD': 'SIM',
        'PPD_Civil': '',
        'Justificativa_Chefe': 'O CURSO IRÁ CONTRIBUIR SIGNIFICATIVAMENTE PARA O APERFEIÇOAMENTO PROFISSIONAL E ATUALIZAÇÃO TÉCNICA DO INDICADO.',
        'Nome_Chefe_COP': 'LEONARDO REZENDE ALVES',
        'Posto_Chefe_COP': 'MAJ AV',
        'Setor_Chefe_COP': 'COP',
        'Nome_Responsavel_DACTA': 'MARCELO MOREIRA DE ANDRADE',
        'Posto_Responsavel_DACTA': 'TEN CEL QOECTA',
        'Setor_Responsavel_DACTA': 'DO',
    }
    
    filler = FICWordFiller()
    output_path = 'FIC_preenchido_exato.docx'
    doc_buffer = filler.preencher_fic(dados_teste, output_path)
    print(f"FIC gerado: {output_path}")
