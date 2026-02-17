from docx import Document

doc = Document('FIC_preenchido_teste.docx')

print('=== NOME COMPLETO COM SUBLINHADO (detalhado) ===')
row = doc.tables[0].rows[7]
cell = row.cells[0]
print(f'Texto completo: {cell.text.strip()}')
print()
for k, p in enumerate(cell.paragraphs):
    print(f'Parágrafo {k}:')
    for run in p.runs:
        underline = 'SUBLINHADO' if run.underline else 'normal'
        print(f'  - "{run.text}" ({underline})')

print()
print('=== CURSO ANTERIOR ===')
row = doc.tables[0].rows[18]
for j, cell in enumerate(row.cells[8:13]):
    if cell.text.strip():
        print(f'C{j+8}: {cell.text.strip()}')
